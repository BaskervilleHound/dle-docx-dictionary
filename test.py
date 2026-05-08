"""Build a formatted Spanish dictionary DOCX from DLE lookups.

The script can look up a single word or read many words from a text file. Found
entries are written into a DOCX with letter headings and Spanish alphabetical
ordering. Words already present in the target DOCX are skipped.
"""

from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from playwright.sync_api import Error, TimeoutError as PlaywrightTimeoutError, sync_playwright


# DLE page selectors.
SEARCH_URL = 'https://dle.rae.es/beta'
SEARCH_INPUT_SELECTORS = [
    'input.js-autocomplete-input',
    'input.c-input--search',
    'input[placeholder="Escriba aquí la palabra"]',
]
SUBMIT_BUTTON_SELECTOR = 'button[type=submit]'
FIRST_DEFINITIONS_SELECTOR = 'ol.c-definitions'
DEFINITIONS_ITEM_SELECTOR = '> li'
ETYMOLOGY_SELECTOR = 'div.c-text-intro'

# Output defaults and dictionary ordering.
DEFAULT_OUTPUT_PATH = 'diccionario.docx'
SPANISH_ALPHABET = 'abcdefghijklmnñopqrstuvwxyz'
SPANISH_ORDER = {letter: index for index, letter in enumerate(SPANISH_ALPHABET)}

# Text cleanup.
GRAMMAR_PREFIX_RE = re.compile(r'^((?:[a-z]\.|[a-z]{2,}\.)\s+)+')
WHITESPACE_RE = re.compile(r'\s+')
WORDS_SPLIT_RE = re.compile(r'[\n,]+')


def get_definitions(word: str) -> tuple[list[str], str]:
    """Fetch the first DLE entry definitions and etymology for a word."""
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        page = browser.new_page(
            user_agent=(
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/125.0.0.0 Safari/537.36'
            )
        )
        page.goto(SEARCH_URL, timeout=60000)

        search_selector = next(
            (selector for selector in SEARCH_INPUT_SELECTORS if page.locator(selector).count()),
            None,
        )
        if search_selector is None:
            raise RuntimeError('Unable to locate the search input on the DLE page.')

        page.fill(search_selector, word)
        page.click(SUBMIT_BUTTON_SELECTOR)

        try:
            page.wait_for_selector(FIRST_DEFINITIONS_SELECTOR, timeout=30000)
        except PlaywrightTimeoutError:
            return [], ''

        first_ol = page.locator(FIRST_DEFINITIONS_SELECTOR).first
        definition_items = first_ol.locator(DEFINITIONS_ITEM_SELECTOR).all()
        definitions = [
            item.locator('div.c-definitions__item > div:not(.c-definitions__item-footer)')
            .first.inner_text()
            .strip()
            for item in definition_items
        ]

        etymology_locator = page.locator(ETYMOLOGY_SELECTOR)
        etymology = etymology_locator.first.inner_text().strip() if etymology_locator.count() else ''

        return definitions, etymology


def read_words_file(words_file: str) -> list[str]:
    """Read comma-separated or newline-separated words from a UTF-8 text file."""
    text = Path(words_file).read_text(encoding='utf-8')
    return [word.strip() for word in WORDS_SPLIT_RE.split(text) if word.strip()]


def normalize_entry_text(text: str) -> str:
    """Normalize DLE text so Word does not wrap non-breaking spaces strangely."""
    return WHITESPACE_RE.sub(' ', text.replace('\xa0', ' ')).strip()


def format_entry_word(word: str) -> str:
    """Keep the entry word lowercase at the first character."""
    return word[:1].lower() + word[1:]


def strip_accents(text: str) -> str:
    """Remove accents for sorting while keeping ñ as its own Spanish letter."""
    text = text.lower().replace('ñ', '\0')
    normalized = unicodedata.normalize('NFD', text)
    stripped = ''.join(char for char in normalized if unicodedata.category(char) != 'Mn')
    return stripped.replace('\0', 'ñ')


def spanish_sort_key(text: str) -> tuple[int, ...]:
    """Return a tuple suitable for Spanish alphabetical ordering."""
    normalized = strip_accents(text)
    return tuple(SPANISH_ORDER.get(char, len(SPANISH_ALPHABET) + ord(char)) for char in normalized)


def entry_key(word: str) -> str:
    """Normalize an entry word for duplicate detection."""
    return strip_accents(format_entry_word(word).strip())


def entry_letter(word: str) -> str:
    """Return the uppercase letter heading for a word."""
    first_char = strip_accents(format_entry_word(word)).strip()[:1]
    return first_char.upper() if first_char else ''


def is_letter_heading(paragraph) -> bool:
    """Detect one-letter section headings such as A, B, C, Ñ."""
    text = paragraph.text.strip()
    return len(text) == 1 and strip_accents(text.lower()) in SPANISH_ORDER


def entry_word_from_paragraph(paragraph) -> str:
    """Extract the dictionary headword from an entry paragraph."""
    text = paragraph.text.strip()
    return text.split('.', 1)[0].strip()


def existing_entry_words(document: Document) -> set[str]:
    """Collect normalized headwords already present in a DOCX."""
    words = set()
    for paragraph in document.paragraphs:
        if is_letter_heading(paragraph):
            continue

        word = entry_word_from_paragraph(paragraph)
        if word:
            words.add(entry_key(word))

    return words


def word_exists_in_docx(word: str, output_path: str) -> bool:
    """Check whether a word is already present in a target DOCX."""
    path = Path(output_path)
    if not path.exists():
        return False

    return entry_key(word) in existing_entry_words(Document(path))


def strip_repeated_grammar_prefix(definition: str, first_prefix: str | None) -> str:
    """Remove repeated grammar labels after the first definition."""
    if first_prefix and definition.startswith(first_prefix):
        return definition[len(first_prefix):].lstrip()
    return definition


def add_letter_heading(document: Document, letter: str):
    """Add a centered, bold letter heading to the document."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(letter)
    run.bold = True
    run.font.size = Pt(14)

    return paragraph


def add_entry_paragraph(document: Document, word: str, definitions: list[str], etymology: str):
    """Add a formatted dictionary entry paragraph at the document end."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    definitions = [normalize_entry_text(definition) for definition in definitions]
    etymology = normalize_entry_text(etymology)

    word_run = paragraph.add_run(f'{format_entry_word(word)}. ')
    word_run.bold = True
    word_run.font.size = Pt(10)

    if etymology:
        opening_parenthesis_run = paragraph.add_run('(')
        opening_parenthesis_run.font.size = Pt(10)

        etymology_run = paragraph.add_run(etymology)
        etymology_run.font.size = Pt(9)

        closing_parenthesis_run = paragraph.add_run(') ')
        closing_parenthesis_run.font.size = Pt(9)

    first_prefix_match = GRAMMAR_PREFIX_RE.match(definitions[0])
    first_prefix = first_prefix_match.group(0) if first_prefix_match else None

    first_definition_run = paragraph.add_run(definitions[0])
    first_definition_run.font.size = Pt(10)

    for index, definition in enumerate(definitions[1:], start=2):
        separator_run = paragraph.add_run(' || ')
        separator_run.font.size = Pt(10)

        number_run = paragraph.add_run(f'{index}.')
        number_run.bold = True
        number_run.font.size = Pt(10)

        definition_run = paragraph.add_run(f' {strip_repeated_grammar_prefix(definition, first_prefix)}')
        definition_run.font.size = Pt(10)

    return paragraph


def move_paragraph_before(paragraph, target) -> None:
    """Move a paragraph before another paragraph using python-docx XML nodes."""
    target._p.addprevious(paragraph._p)


def move_paragraph_after(paragraph, target) -> None:
    """Move a paragraph after another paragraph using python-docx XML nodes."""
    target._p.addnext(paragraph._p)


def delete_paragraph(paragraph) -> None:
    """Delete a paragraph from the document."""
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def remove_empty_paragraphs(document: Document) -> None:
    """Remove blank paragraphs so entries remain compact."""
    for paragraph in list(document.paragraphs):
        if not paragraph.text.strip():
            delete_paragraph(paragraph)


def find_letter_heading(document: Document, letter: str):
    """Find an existing letter heading paragraph."""
    for paragraph in document.paragraphs:
        if is_letter_heading(paragraph) and paragraph.text.strip().upper() == letter:
            return paragraph
    return None


def paragraph_index(document: Document, target) -> int:
    """Find a paragraph by its XML node instead of wrapper identity."""
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError('Paragraph is not in the document.')


def insert_heading_in_order(document: Document, heading) -> None:
    """Move a new heading into Spanish alphabetical position."""
    heading_key = SPANISH_ORDER[strip_accents(heading.text.lower())]

    for paragraph in document.paragraphs:
        if not is_letter_heading(paragraph):
            continue

        paragraph_key = SPANISH_ORDER[strip_accents(paragraph.text.lower())]
        if paragraph_key > heading_key:
            move_paragraph_before(heading, paragraph)
            return


def insert_entry_in_order(document: Document, entry_paragraph, word: str) -> None:
    """Insert an entry under its heading and sort that heading section."""
    remove_empty_paragraphs(document)

    letter = entry_letter(word)
    heading = find_letter_heading(document, letter)
    if heading is None:
        heading = add_letter_heading(document, letter)
        insert_heading_in_order(document, heading)

    move_paragraph_after(entry_paragraph, heading)

    paragraphs = document.paragraphs
    heading_index = paragraph_index(document, heading)
    section_entries = []

    for paragraph in paragraphs[heading_index + 1:]:
        if is_letter_heading(paragraph):
            break

        current_word = entry_word_from_paragraph(paragraph)
        if current_word:
            section_entries.append(paragraph)

    section_entries.sort(key=lambda paragraph: spanish_sort_key(entry_word_from_paragraph(paragraph)))
    for paragraph in reversed(section_entries):
        move_paragraph_after(paragraph, heading)


def save_to_docx(word: str, definitions: list[str], etymology: str, output_path: str) -> None:
    """Create or update a DOCX with one formatted entry."""
    path = Path(output_path)
    document = Document(path) if path.exists() else Document()
    entry_paragraph = add_entry_paragraph(document, word, definitions, etymology)
    insert_entry_in_order(document, entry_paragraph, word)
    document.save(path)


def build_parser() -> argparse.ArgumentParser:
    """Create the command-line parser."""
    parser = argparse.ArgumentParser(
        description='Scrape the DLE online dictionary and write entries to a DOCX.'
    )
    parser.add_argument('word', nargs='?', help='Word to look up')
    parser.add_argument(
        '-w',
        '--words-file',
        nargs='?',
        const='words.txt',
        help='Read words from a text file (default when used without a value: words.txt)',
    )
    parser.add_argument(
        '-o',
        '--output',
        default=DEFAULT_OUTPUT_PATH,
        help=f'DOCX file to write results to (default: {DEFAULT_OUTPUT_PATH})',
    )
    return parser


def print_run_summary(not_found_words: list[str], already_existing_words: list[str]) -> None:
    """Print the final status summary for a run."""
    if not_found_words:
        print('Words not found in DLE:')
        for word in not_found_words:
            print(f'- {word}')
    else:
        print('All words were found in DLE.')

    if already_existing_words:
        print('Words already present in the DOCX and skipped:')
        for word in already_existing_words:
            print(f'- {word}')
    else:
        print('No requested words were already present in the DOCX.')


def main() -> int:
    """Run the command-line workflow."""
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')

    parser = build_parser()
    args = parser.parse_args()

    if args.words_file:
        words = read_words_file(args.words_file)
    elif args.word:
        words = [args.word]
    else:
        parser.error('provide a word or use --words-file')

    not_found_words = []
    already_existing_words = []
    existing_words = existing_entry_words(Document(args.output)) if Path(args.output).exists() else set()

    for word in words:
        if entry_key(word) in existing_words:
            already_existing_words.append(word)
            continue

        try:
            definitions, etymology = get_definitions(word)
        except Error as exc:
            print(f'Playwright error for "{word}": {exc}')
            continue
        except Exception as exc:
            print(f'Lookup failed for "{word}": {exc}')
            continue

        if not definitions:
            not_found_words.append(word)
            print(f'No definition found for "{word}" in DLE.')
            continue

        save_to_docx(word, definitions, etymology, args.output)
        existing_words.add(entry_key(word))

        print(f'Definitions for "{word}":')
        print('Etimología:')
        print(etymology)
        print('Definiciones:')
        for index, definition in enumerate(definitions, start=1):
            print(f'{index}. {definition}')
        print(f'Output saved to {args.output}')

    print_run_summary(not_found_words, already_existing_words)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
